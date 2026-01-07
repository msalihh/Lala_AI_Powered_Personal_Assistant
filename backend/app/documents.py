"""
Document management utilities - Text extraction from PDF/DOCX/TXT/IMAGE.
"""
import re
import logging
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

# Text extraction limits
MAX_TEXT_LENGTH = 200000  # 200k characters
MAX_PDF_PAGES = 200  # Max pages to extract from PDF
MAX_DOCX_PARAGRAPHS = 10000  # Max paragraphs to extract from DOCX


def extract_text_from_pdf(file_content: bytes) -> Tuple[str, bool]:
    """
    Extract text from PDF file using PyMuPDF.
    Returns: (text, was_truncated)
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        page_count = len(doc)
        was_truncated = False
        
        logger.info(f"[PDF_EXTRACT] Starting extraction: pages={page_count}, file_size={len(file_content)} bytes")
        
        # Limit pages
        max_pages = min(page_count, MAX_PDF_PAGES)
        if page_count > MAX_PDF_PAGES:
            was_truncated = True
            logger.warning(f"[PDF_EXTRACT] PDF has {page_count} pages, limiting to {MAX_PDF_PAGES}")
        
        total_chars_extracted = 0
        pages_with_text = 0
        pages_without_text = 0
        
        for i in range(max_pages):
            page = doc[i]
            page_text = page.get_text()
            text_parts.append(page_text)
            
            if page_text.strip():
                pages_with_text += 1
                total_chars_extracted += len(page_text)
            else:
                pages_without_text += 1
                logger.warning(f"[PDF_EXTRACT] Page {i+1} returned empty text (may be scanned/image-based PDF)")
        
        doc.close()
        full_text = "\n".join(text_parts)
        
        logger.info(
            f"[PDF_EXTRACT] Extraction complete: "
            f"total_chars={len(full_text)}, "
            f"pages_with_text={pages_with_text}, "
            f"pages_without_text={pages_without_text}, "
            f"is_empty={len(full_text.strip()) == 0}"
        )
        
        # Limit text length
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]
            was_truncated = True
            logger.warning(f"[PDF_EXTRACT] Text truncated to {MAX_TEXT_LENGTH} characters")
        
        # Warn if extracted text is empty or very short
        if len(full_text.strip()) == 0:
            logger.error(
                f"[PDF_EXTRACT] WARNING: Extracted text is EMPTY! "
                f"This PDF may be scanned/image-based and requires OCR. "
                f"Pages checked: {max_pages}, pages_with_text: {pages_with_text}"
            )
        elif len(full_text.strip()) < 50:
            logger.warning(
                f"[PDF_EXTRACT] WARNING: Extracted text is very short ({len(full_text.strip())} chars). "
                f"This may indicate a scanned PDF or extraction issue."
            )
        
        return full_text, was_truncated
    except Exception as e:
        logger.error(f"[PDF_EXTRACT] Error extracting PDF: {str(e)}", exc_info=True)
        raise ValueError(f"PDF okuma hatası: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> Tuple[str, bool]:
    """
    Extract text from DOCX file using python-docx.
    Returns: (text, was_truncated)
    """
    try:
        from docx import Document
        from io import BytesIO
        
        doc = Document(BytesIO(file_content))
        text_parts = []
        paragraph_count = 0
        was_truncated = False
        
        # Extract paragraphs with limit
        for paragraph in doc.paragraphs:
            if paragraph_count >= MAX_DOCX_PARAGRAPHS:
                was_truncated = True
                break
            text_parts.append(paragraph.text)
            paragraph_count += 1
        
        # Also extract text from tables (limited)
        table_count = 0
        for table in doc.tables:
            if table_count >= 100:  # Limit tables too
                was_truncated = True
                break
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                text_parts.append(" | ".join(row_text))
            table_count += 1
        
        full_text = "\n".join(text_parts)
        
        # Limit text length
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]
            was_truncated = True
        
        return full_text, was_truncated
    except Exception as e:
        raise ValueError(f"DOCX okuma hatası: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> Tuple[str, bool]:
    """
    Extract text from TXT file with UTF-8 fallback to latin-1.
    Returns: (text, was_truncated)
    """
    try:
        # Try UTF-8 first
        text = file_content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            # Fallback to latin-1
            text = file_content.decode("latin-1")
        except Exception as e:
            raise ValueError(f"TXT okuma hatası: {str(e)}")
    
    # Limit text length
    was_truncated = False
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH]
        was_truncated = True
    
    return text, was_truncated


def extract_text_from_image(file_content: bytes, filename: str) -> Tuple[str, bool]:
    """
    Extract text from image using OCR and vision analysis.
    Returns: (combined_text, was_truncated)
    Combined text includes OCR text + caption for RAG indexing.
    """
    try:
        from app.vision import analyze_image
        
        logger.info(f"[IMAGE_EXTRACT] Starting image analysis: filename={filename}, size={len(file_content)} bytes")
        
        # Run image analysis (OCR + Vision)
        analysis = analyze_image(file_content, filename)
        
        # Combine OCR text and caption for RAG indexing
        text_parts = []
        if analysis.get("ocr_text"):
            text_parts.append(analysis["ocr_text"])
        if analysis.get("caption"):
            text_parts.append(f"Fotoğraf açıklaması: {analysis['caption']}")
        
        combined_text = "\n\n".join(text_parts) if text_parts else ""
        
        # Limit text length
        was_truncated = False
        if len(combined_text) > MAX_TEXT_LENGTH:
            combined_text = combined_text[:MAX_TEXT_LENGTH]
            was_truncated = True
        
        logger.info(
            f"[IMAGE_EXTRACT] Extraction complete: "
            f"ocr_text_length={len(analysis.get('ocr_text', ''))}, "
            f"caption_length={len(analysis.get('caption', ''))}, "
            f"combined_length={len(combined_text)}, "
            f"was_truncated={was_truncated}"
        )
        
        return combined_text, was_truncated
        
    except Exception as e:
        logger.error(f"[IMAGE_EXTRACT] Error extracting image: {str(e)}", exc_info=True)
        # Return empty text if analysis fails (system should still work)
        return "", False


def extract_text_from_file(file_content: bytes, mime_type: str, filename: str) -> Tuple[str, bool]:
    """
    Extract text from file based on MIME type.
    Returns: (normalized_text, was_truncated)
    Supports: PDF, DOCX, TXT, IMAGE (jpg/png/webp)
    """
    was_truncated = False
    
    logger.info(f"[TEXT_EXTRACT] Starting extraction: filename={filename}, mime_type={mime_type}, size={len(file_content)} bytes")
    
    # Check for image types first (before strict validation)
    image_mime_types = {
        "image/jpeg",
        "image/jpg",
        "image/png",
        "image/webp"
    }
    
    if mime_type in image_mime_types:
        # Image files: use OCR + vision
        image_extensions = {".jpg", ".jpeg", ".png", ".webp"}
        file_ext = filename.lower()
        if any(file_ext.endswith(ext) for ext in image_extensions):
            raw_text, truncated = extract_text_from_image(file_content, filename)
            was_truncated = was_truncated or truncated
            # Images don't need normalization (already clean from OCR/vision)
            return raw_text, was_truncated
        else:
            raise ValueError(f"Görüntü dosyası geçerli bir uzantıya sahip olmalıdır: {', '.join(image_extensions)}")
    
    # STRICT validation: Only allow specific MIME types for documents
    if mime_type == "application/pdf":
        if not filename.lower().endswith(".pdf"):
            raise ValueError("PDF dosyası .pdf uzantısına sahip olmalıdır")
        raw_text, truncated = extract_text_from_pdf(file_content)
        was_truncated = was_truncated or truncated
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        if not filename.lower().endswith(".docx"):
            raise ValueError("DOCX dosyası .docx uzantısına sahip olmalıdır")
        raw_text, truncated = extract_text_from_docx(file_content)
        was_truncated = was_truncated or truncated
    elif mime_type == "text/plain":
        if not filename.lower().endswith(".txt"):
            raise ValueError("TXT dosyası .txt uzantısına sahip olmalıdır")
        raw_text, truncated = extract_text_from_txt(file_content)
        was_truncated = was_truncated or truncated
    else:
        raise ValueError(f"Desteklenmeyen dosya tipi: {mime_type}. İzin verilen: application/pdf, application/vnd.openxmlformats-officedocument.wordprocessingml.document, text/plain, image/jpeg, image/png, image/webp")
    
    logger.info(f"[TEXT_EXTRACT] Raw text extracted: length={len(raw_text)}, has_content={bool(raw_text.strip())}")
    
    # Normalize text: remove excessive whitespace, normalize line breaks
    normalized = normalize_text(raw_text)
    
    logger.info(
        f"[TEXT_EXTRACT] Normalization complete: "
        f"raw_length={len(raw_text)}, "
        f"normalized_length={len(normalized)}, "
        f"normalized_has_content={bool(normalized.strip())}, "
        f"was_truncated={was_truncated}"
    )
    
    # Final check: warn if normalized text is empty
    if not normalized.strip():
        logger.error(
            f"[TEXT_EXTRACT] CRITICAL: Normalized text is EMPTY after extraction! "
            f"filename={filename}, mime_type={mime_type}, "
            f"raw_text_length={len(raw_text)}, normalized_length={len(normalized)}"
        )
    
    return normalized, was_truncated


def normalize_text(text: str) -> str:
    """
    Normalize text: remove excessive whitespace, normalize line breaks.
    """
    # Replace multiple spaces with single space
    text = re.sub(r" +", " ", text)
    # Replace multiple newlines with max 2 newlines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split("\n")]
    # Remove empty lines at start/end
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop(-1)
    return "\n".join(lines)

